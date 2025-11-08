import streamlit as st
import pandas as pd
import numpy as np
from sklearn.ensemble import RandomForestRegressor
from sklearn.linear_model import LinearRegression
from sklearn.multioutput import MultiOutputRegressor
from sklearn.metrics import r2_score, mean_squared_error
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches

st.set_page_config(page_title="Ore Abrasivity & Breakability Predictor", layout="wide")

st.title("Ore Abrasivity & Breakability Predictor (ML + Regression)")
st.write("Upload your ore dataset CSV and get predictions for Abrasivity & Breakability, along with best conditioner selection.")

# -------------------------------
# CSV Upload
# -------------------------------
uploaded_file = st.file_uploader("Upload your CSV file", type="csv")

if uploaded_file is not None:
    df = pd.read_csv(uploaded_file)
    st.success(f"Loaded {df.shape[0]} rows and {df.shape[1]} columns")

    # -------------------------------
    # User input
    # -------------------------------
    moisture = st.slider("Select moisture % for predictions", min_value=5, max_value=25, value=10, step=5)

    # -------------------------------
    # Model training function
    # -------------------------------
    def train_models(df):
        X = df[['Quartz %', 'Bentonite %', 'Moisture %', 'Conditioner']]
        y = df[['Abrasivity', 'Breakability']]
        X_encoded = pd.get_dummies(X, columns=['Conditioner'])

        # Random Forest
        rf_model = MultiOutputRegressor(RandomForestRegressor(n_estimators=200, random_state=42))
        rf_model.fit(X_encoded, y)

        # Linear Regression
        lr_model = MultiOutputRegressor(LinearRegression())
        lr_model.fit(X_encoded, y)

        return rf_model, lr_model, X_encoded, y

    # -------------------------------
    # Process uploaded data
    # -------------------------------
    def process_ore_data(df, moisture):
        rf_model, lr_model, X_encoded_full, y_full = train_models(df)

        # Predictions
        y_pred_rf = rf_model.predict(X_encoded_full)
        y_pred_lr = lr_model.predict(X_encoded_full)

        # Metrics
        r2_abrasivity_rf = r2_score(y_full['Abrasivity'], y_pred_rf[:,0])
        r2_breakability_rf = r2_score(y_full['Breakability'], y_pred_rf[:,1])
        rmse_abrasivity_rf = np.sqrt(mean_squared_error(y_full['Abrasivity'], y_pred_rf[:,0]))
        rmse_breakability_rf = np.sqrt(mean_squared_error(y_full['Breakability'], y_pred_rf[:,1]))

        r2_abrasivity_lr = r2_score(y_full['Abrasivity'], y_pred_lr[:,0])
        r2_breakability_lr = r2_score(y_full['Breakability'], y_pred_lr[:,1])
        rmse_abrasivity_lr = np.sqrt(mean_squared_error(y_full['Abrasivity'], y_pred_lr[:,0]))
        rmse_breakability_lr = np.sqrt(mean_squared_error(y_full['Breakability'], y_pred_lr[:,1]))

        # Batch best-conditioner prediction
        ratios = range(0, 101, 10)
        conditioners = ['None', 'A', 'B', 'C']
        best_results = []
        feature_cols = X_encoded_full.columns

        for quartz in ratios:
            bentonite = 100 - quartz
            results = []
            for cond in conditioners:
                input_df = pd.DataFrame([[quartz, bentonite, moisture, cond]],
                                        columns=['Quartz %', 'Bentonite %', 'Moisture %', 'Conditioner'])
                input_encoded = pd.get_dummies(input_df, columns=['Conditioner'])
                for col in feature_cols:
                    if col not in input_encoded:
                        input_encoded[col] = 0
                input_encoded = input_encoded[feature_cols]

                rf_pred = rf_model.predict(input_encoded)[0]
                lr_pred = lr_model.predict(input_encoded)[0]

                results.append({
                    'Conditioner': cond,
                    'RF_Abrasivity': rf_pred[0],
                    'RF_Breakability': rf_pred[1],
                    'LR_Abrasivity': lr_pred[0],
                    'LR_Breakability': lr_pred[1],
                    'Score': rf_pred[1] - rf_pred[0]
                })

            results_df = pd.DataFrame(results)
            best = results_df.loc[results_df['Score'].idxmax()]
            best_results.append({
                'Quartz %': quartz,
                'Bentonite %': bentonite,
                'Moisture %': moisture,
                'Best Conditioner': best['Conditioner'],
                'RF_Abrasivity': best['RF_Abrasivity'],
                'RF_Breakability': best['RF_Breakability'],
                'LR_Abrasivity': best['LR_Abrasivity'],
                'LR_Breakability': best['LR_Breakability']
            })

        best_results_df = pd.DataFrame(best_results)
        return (rf_model, lr_model, X_encoded_full, y_full,
                y_pred_rf, y_pred_lr,
                r2_abrasivity_rf, r2_breakability_rf,
                rmse_abrasivity_rf, rmse_breakability_rf,
                r2_abrasivity_lr, r2_breakability_lr,
                rmse_abrasivity_lr, rmse_breakability_lr,
                best_results_df)

    with st.spinner("Processing data and training models..."):
        results = process_ore_data(df, moisture)
        (rf_model, lr_model, X_encoded_full, y_full,
         y_pred_rf, y_pred_lr,
         r2_abrasivity_rf, r2_breakability_rf,
         rmse_abrasivity_rf, rmse_breakability_rf,
         r2_abrasivity_lr, r2_breakability_lr,
         rmse_abrasivity_lr, rmse_breakability_lr,
         best_results_df) = results

    st.success("Processing complete!")

    st.subheader("Best Conditioner Predictions")
    st.dataframe(best_results_df.style.format({
        "RF_Abrasivity": "{:.3f}",
        "RF_Breakability": "{:.3f}",
        "LR_Abrasivity": "{:.3f}",
        "LR_Breakability": "{:.3f}"
    }))

    # -------------------------------
    # Word report generation
    # -------------------------------
    def generate_word(df, moisture,
                      r2_abrasivity_rf, r2_breakability_rf,
                      rmse_abrasivity_rf, rmse_breakability_rf,
                      r2_abrasivity_lr, r2_breakability_lr,
                      rmse_abrasivity_lr, rmse_breakability_lr):
        doc = Document()
        doc.add_heading("Ore Prediction Report", 0)

        # ML Model
        doc.add_heading("Machine Learning Model", level=1)
        doc.add_paragraph("Random Forest Regressor (MultiOutput)")
        doc.add_paragraph("Preprocessing: One-hot encoding of Conditioner")

        # Regression Model
        doc.add_heading("Regression Model", level=1)
        doc.add_paragraph("Linear Regression (MultiOutput)")
        doc.add_paragraph("Preprocessing: One-hot encoding of Conditioner")

        # Metrics
        doc.add_heading("Model Performance Metrics", level=1)
        doc.add_paragraph(f"Moisture used: {moisture}%")
        doc.add_paragraph("Random Forest Performance:")
        doc.add_paragraph(f"Abrasivity - R²: {r2_abrasivity_rf:.3f}, RMSE: {rmse_abrasivity_rf:.3f}")
        doc.add_paragraph(f"Breakability - R²: {r2_breakability_rf:.3f}, RMSE: {rmse_breakability_rf:.3f}")
        doc.add_paragraph("Linear Regression Performance:")
        doc.add_paragraph(f"Abrasivity - R²: {r2_abrasivity_lr:.3f}, RMSE: {rmse_abrasivity_lr:.3f}")
        doc.add_paragraph(f"Breakability - R²: {r2_breakability_lr:.3f}, RMSE: {rmse_breakability_lr:.3f}")

        # Table
        doc.add_heading("Best Conditioner for Each Ore Mixture", level=1)
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = col
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(df.columns):
                row_cells[i].text = str(round(row[col],3)) if isinstance(row[col], float) else str(row[col])

        # Graphs
        doc.add_page_break()
        doc.add_heading("Predicted vs Actual Graphs", level=1)

        fig, ax = plt.subplots(2,2, figsize=(10,8))
        # RF Abrasivity
        ax[0,0].scatter(y_full['Abrasivity'], y_pred_rf[:,0], color='blue')
        ax[0,0].plot([0,1],[0,1], 'r--')
        ax[0,0].set_xlabel("Actual Abrasivity")
        ax[0,0].set_ylabel("Predicted Abrasivity")
        ax[0,0].set_title("RF: Abrasivity Predicted vs Actual")

        # RF Breakability
        ax[0,1].scatter(y_full['Breakability'], y_pred_rf[:,1], color='green')
        ax[0,1].plot([0,1],[0,1], 'r--')
        ax[0,1].set_xlabel("Actual Breakability")
        ax[0,1].set_ylabel("Predicted Breakability")
        ax[0,1].set_title("RF: Breakability Predicted vs Actual")

        # LR Abrasivity
        ax[1,0].scatter(y_full['Abrasivity'], y_pred_lr[:,0], color='blue')
        ax[1,0].plot([0,1],[0,1], 'r--')
        ax[1,0].set_xlabel("Actual Abrasivity")
        ax[1,0].set_ylabel("Predicted Abrasivity")
        ax[1,0].set_title("LR: Abrasivity Predicted vs Actual")

        # LR Breakability
        ax[1,1].scatter(y_full['Breakability'], y_pred_lr[:,1], color='green')
        ax[1,1].plot([0,1],[0,1], 'r--')
        ax[1,1].set_xlabel("Actual Breakability")
        ax[1,1].set_ylabel("Predicted Breakability")
        ax[1,1].set_title("LR: Breakability Predicted vs Actual")

        img_stream = BytesIO()
        plt.tight_layout()
        plt.savefig(img_stream, format='png')
        plt.close(fig)
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(6))

        return doc

    if st.button("Download Word Report"):
        doc = generate_word(best_results_df, moisture,
                            r2_abrasivity_rf, r2_breakability_rf,
                            rmse_abrasivity_rf, rmse_breakability_rf,
                            r2_abrasivity_lr, r2_breakability_lr,
                            rmse_abrasivity_lr, rmse_breakability_lr)
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        st.download_button(
            label="Download Word Report",
            data=file_stream,
            file_name="Ore_Prediction_Comparison_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
